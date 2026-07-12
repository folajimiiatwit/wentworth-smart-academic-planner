"""
Purpose:
Extracts completed courses from uploaded transcript files.

Main responsibilities:
- Extract text from PDF, DOCX, and TXT transcripts (Currently TXT is not working and won't be made avaliable to upload)
- Ignore courses listed under COURSE(S) IN PROGRESS
- Ignore failed or withdrawn courses
- Identify completed course codes
- Separate required courses from custom or transfer courses
"""
import re
from io import BytesIO
from docx import Document
import pdfplumber

PASSING_GRADES = {
    "A","A-",
    "B+","B","B-",
    "C+","C","C-",
    "D+","D","S","TR",
}

FAILING_GRADES = {
    "F", "W", "D-",

}

SEMESTER_WORDS = {
    "FALL", "SPRING", "SUMMER", "WINTER", "TERM", "YEAR",
}

EXCLUDED_SUBJECTS = {
    "COOP"
}

NON_COURSE_WORDS = {
    "GPA", "UG", "CEU", "WEB", "TOP", "THE", "AND", "FOR", "NOT",
    "TOTAL", "TOTALS", "CURRENT", "CUMULATIVE", "OVERALL", "COOP",
}

STOP_MARKERS = [
    "COURSE(S) IN PROGRESS",
    
]

GRADE_PATTERN = (
    r"(?<![A-Z0-9])"
    r"(A-|A|B\+|B-|B|C\+|C-|C|D\+|D-|D|TR|P|S|F|W|WF|WU|NC|I)"
    r"(?![A-Z0-9])"
)

def extract_text_from_docx(file_bytes):
    """
    Extract transcript text from a DOCX file.

    Both paragraph text and table cell text are included because transcripts may store
    course information in either format.

    Args:
        file_bytes (bytes): Raw DOCX file content.

    Returns:
        str: Extracted transcript text.
    """
    document = Document(BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return "\n".join(parts)


def extract_text_from_pdf(file_bytes):
    """
    Extract transcript text from a PDF file using pdfplumber.

    Args:
        file_bytes (bytes): Raw PDF file content.

    Returns:
        str: Extracted transcript text from all readable pages.
    """

    parts = []

    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            words = page.extract_words(
                use_text_flow=True,
                keep_blank_chars=False
            )

            if not words:
                continue

            visual_lines = {}
            tolerance = 3

            for word in words:
                row_position = round(word["top"] / tolerance) * tolerance
                visual_lines.setdefault(row_position, []).append(word)

            for row_position in sorted(visual_lines):
                row_words = sorted(
                    visual_lines[row_position],
                    key=lambda item: item["x0"]
                )

                line = " ".join(
                    word["text"] for word in row_words
                ).strip()

                if line:
                    parts.append(line)

    return "\n".join(parts)


def extract_text_from_txt(file_bytes):
    """
    Extract transcript text from a plain-text file.

    Args:
        file_bytes (bytes): Raw TXT file content.

    Returns:
        str: Decoded transcript text.
    """
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_file(filename, file_bytes):
    """
    Select the correct transcript text extractor based on file extension.

    Supported extensions include `.docx`, `.pdf`, and `.txt`. Unknown file types are
    processed as plain text.

    Args:
        filename (str): Uploaded transcript filename.
        file_bytes (bytes): Raw uploaded file content.

    Returns:
        str: Extracted transcript text.
    """
    filename = filename.lower()

    if filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    if filename.endswith(".txt"):
        return extract_text_from_txt(file_bytes)

    return extract_text_from_txt(file_bytes)


def remove_in_progress_section(text):
    """
    Remove transcript sections that list courses currently in progress.

    Courses after the configured stop marker should not count as completed courses.

    Args:
        text (str): Full transcript text.

    Returns:
        str: Transcript text before any in-progress section.
    """
    match = re.search(
        r"COURSE\s*\(S\)\s*IN\s*PROGRESS|COURSES?\s*IN\s*PROGRESS",
        text,
        flags=re.IGNORECASE
    )

    if match:
        return text[:match.start()]

    return text


def normalize_lines(text):
    """
    Convert transcript text into a clean list of non-empty lines.

    Args:
        text (str): Raw transcript text.

    Returns:
        list[str]: Trimmed non-empty transcript lines.
    """
    lines = []

    for line in text.splitlines():
        cleaned = line.strip()
        if cleaned:
            lines.append(cleaned)

    return lines


def is_valid_subject(subject):
    """
    Check whether a detected subject prefix looks like a real course subject.

    Semester words and common non-course words are rejected to reduce false matches.

    Args:
        subject (str): Candidate course subject prefix.

    Returns:
        bool: True if the subject appears valid.
    """
    subject = subject.strip().upper()

    if not re.fullmatch(r"[A-Z]{3,5}", subject):
        return False

    if subject in SEMESTER_WORDS:
        return False

    if subject in NON_COURSE_WORDS:
        return False
    
    if subject in EXCLUDED_SUBJECTS:
        return False

    return True


def looks_like_course_number(value):
    """
    Check whether a value looks like a four-digit course number.

    Args:
        value (str): Candidate course number.

    Returns:
        bool: True if the value is exactly four digits.
    """
    return bool(re.fullmatch(r"\d{4}", value.strip()))


def is_valid_course_pair(subject, number):
    """
    Validate a subject and course-number pair before treating it as a course.

    This prevents transcript headings such as `FALL 2025` from being mistaken for
    course codes.

    Args:
        subject (str): Candidate course subject.
        number (str): Candidate four-digit course number.

    Returns:
        bool: True if the pair appears to represent a course.
    """
    subject = subject.strip().upper()
    number = number.strip()

    if not is_valid_subject(subject):
        return False

    if not looks_like_course_number(number):
        return False

    # Extra safety: do not treat semester years as course numbers.
    if subject in SEMESTER_WORDS and number.startswith("20"):
        return False

    return True

def combine_course_rows(lines):
    """
    Combine wrapped transcript lines into complete course records.

    A new record starts whenever a line contains a subject and four-digit
    course number. Following wrapped lines are added until the next course
    begins.
    """
    rows = []
    current_row = ""

    for line in lines:
        starts_course = bool(
            re.search(r"\b[A-Z]{3,5}\s+\d{4}\b", line)
        )

        if starts_course:
            if current_row:
                rows.append(current_row.strip())

            current_row = line
        elif current_row:
            current_row += " " + line

    if current_row:
        rows.append(current_row.strip())

    return rows

def extract_completed_courses_from_text(text):
    """
    Extract completed course codes from transcript text.

    Only courses with passing grades are included. Courses in progress, failed
    courses, withdrawn courses, and semester headings are ignored.

    Args:
        text (str): Extracted transcript text.

    Returns:
        list[str]: Sorted completed course codes.
    """
    completed = set()
    safe_text = remove_in_progress_section(text)
    lines = normalize_lines(safe_text.upper())
    course_rows = combine_course_rows(lines)

    for row in course_rows:
        match = re.search(
        r"\b([A-Z]{3,5})\s+(\d{4})\b",
        row
        )

        if not match:
            continue

        subject, number = match.groups()

        if not is_valid_course_pair(subject, number):
          continue

        grade_match = re.search(
            r"(A-|A|B\+|B-|B|C\+|C-|C|D\+|D-|D|TR|P|S|F|W|WF|WU|NC|I)"
            r"\s+\d+\.\d{3}\s+\d+\.\d{2}\b",
            row
        )

        if not grade_match:
            continue

        grade = grade_match.group(1).upper()

        if grade in PASSING_GRADES:
            completed.add(f"{subject}{number}")

    return sorted(completed)


def split_required_and_custom(extracted_codes, required_course_codes):
    """
    Separate extracted transcript courses into required and custom completed courses.

    Args:
        extracted_codes (list[str]): Course codes found in the transcript.
        required_course_codes (list[str]): Required course codes for the program.

    Returns:
        tuple: Completed required course codes and custom completed course records.
    """
    required_set = set(str(code).strip().upper() for code in required_course_codes)

    completed_required = []
    custom_completed = []

    for code in extracted_codes:
        code = str(code).strip().upper()

        if code in required_set:
            completed_required.append(code)
        else:
            subject_match = re.match(r"([A-Z]+)(\d+)", code)

            if subject_match:
                number = subject_match.group(2)
            else:
                number = ""

            custom_completed.append({
                "course_code": code,
                "course_number": number,
                "title": "Imported from transcript"
            })

    return completed_required, custom_completed
